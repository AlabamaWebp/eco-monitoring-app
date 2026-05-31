# -*- coding: utf-8 -*-
"""Generate the course report DOCX from course_report.md.

The script reads the editable Markdown file and produces a Word document with
formatting close to common Russian coursework requirements: A4, Times New Roman,
14 pt body text, 1.5 line spacing, justified paragraphs, ГОСТ-like margins,
tables with borders, centered figure captions, and page numbering.

If course_report.md is missing, the script writes the default prepared text
embedded below. If the Markdown file already exists, it is used as the source and
is not overwritten.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "course_report.md"
DOCX_PATH = BASE_DIR / "course_report.docx"


DEFAULT_MARKDOWN = r"""# Введение

Экологический мониторинг связан с регулярным сбором числовых показателей состояния окружающей среды. Для анализа таких данных необходимо хранить не только сами значения, но и сведения о месте наблюдения, типе показателя, единице измерения, времени измерения и происхождении записи. Если данные остаются в отдельных CSV-файлах, возникает несколько практических проблем: файлы трудно сравнивать между собой, сложно быстро фильтровать измерения по периоду и полигону, неудобно строить графики и контролировать историю загрузок.

Тема курсовой работы является актуальной, поскольку сетевые базы данных позволяют централизованно хранить экологические измерения и предоставлять к ним доступ через веб-приложение. В разработанном проекте база данных размещается в MySQL, а пользователь работает с ней через frontend-интерфейс. Это позволяет выполнять импорт CSV-файлов, просматривать измерения, применять фильтры, строить графики, сравнивать полигоны, а также вручную добавлять, редактировать и удалять отдельные записи.

Цель курсовой работы — разработать базу данных и веб-приложение для хранения и анализа экологических измерений, собираемых на разных полигонах.

Для достижения цели были поставлены следующие задачи:

- изучить предметную область экологических измерений;
- определить формат исходных CSV-данных;
- спроектировать нормализованную структуру базы данных MySQL;
- реализовать таблицы, первичные и внешние ключи, индексы и ограничения;
- обеспечить импорт широкого CSV-формата в нормализованную таблицу измерений;
- реализовать просмотр, фильтрацию и построение графиков;
- реализовать ручное добавление, редактирование и удаление измерений;
- реализовать управление справочниками полигонов, типов показателей и единиц измерения;
- обеспечить защиту связанных справочных данных от некорректного удаления.

Объектом работы являются экологические измерения, собираемые на полигонах. Предметом работы является структура сетевой базы данных и программные механизмы доступа к ней через веб-приложение.

В проекте используются MySQL 8, FastAPI, SQLAlchemy ORM, Alembic, Angular, TypeScript, Apache ECharts и Docker Compose. Отчет включает анализ предметной области, проектирование базы данных в MySQL, описание серверной работы с данными, пользовательский интерфейс, SQL-запросы и итоговые выводы.

# 1. Анализ предметной области и постановка задачи

## 1.1. Описание предметной области экологических измерений

В предметной области рассматривается несколько экологических полигонов. На каждом полигоне регулярно собираются значения показателей окружающей среды: CO2, влажность, температура, давление, освещенность, уровень шума и другие величины, если они заведены в справочнике типов показателей.

В реализованной модели тип показателя описывает измеряемую величину, а не физическое устройство. Поэтому в базе данных нет отдельной сущности физического датчика и нет привязки устройств к полигонам. Факт наблюдения определяется полигоном, типом показателя, единицей измерения, датой и временем, а также числовым значением.

Данные могут попадать в систему двумя способами. Первый способ — загрузка CSV-файла, подготовленного человеком. При загрузке пользователь выбирает полигон и указывает фамилию загрузившего. Второй способ — ручное добавление отдельного измерения через форму на странице измерений. Оба способа создают записи в одной таблице measurements, поэтому такие данные участвуют в общих фильтрах, таблицах и графиках.

## 1.2. Формат исходных данных

Исходные CSV-файлы имеют широкий формат. Одна строка соответствует одному моменту времени, а отдельные столбцы после даты соответствуют разным показателям.

```csv
Дата,CO2,Влажность,Температура,Давление,Освещённость,Уровень шума
2025-01-01 00:00:00,420,55,21.4,1008,610,38
```

При импорте широкий формат преобразуется в длинный формат. Если строка CSV содержит дату и шесть заполненных показателей, в таблице measurements создается шесть строк. В каждой строке хранится один показатель, его значение, время измерения, полигон, единица измерения и ссылка на файл импорта.

Ручная запись создается без CSV-файла. В этом случае пользователь указывает полигон, тип показателя, время и значение, а единица измерения определяется автоматически по выбранному типу показателя.

## 1.3. Требования к базе данных

База данных должна выполнять следующие функции:

- хранить полигоны экологического мониторинга;
- хранить загрузивших CSV-файлы сотрудников;
- хранить единицы измерения;
- хранить типы измеряемых показателей;
- хранить журнал загрузок CSV-файлов;
- хранить импортированные и ручные измерения в общей таблице;
- допускать NULL в measurements.import_file_id для ручных записей;
- определять единицу измерения ручной записи через выбранный тип показателя;
- поддерживать фильтрацию по полигону, типу показателя, периоду и файлу импорта;
- поддерживать выборки для построения графиков и сравнения полигонов;
- предотвращать удаление справочных записей, если они используются существующими измерениями или загрузками;
- сохранять ссылочную целостность при операциях добавления, редактирования и удаления.

Структура базы данных должна быть нормализована до третьей нормальной формы. Названия полигонов, типов показателей и единиц измерения не должны дублироваться текстом в таблице измерений.

## 1.4. Требования к веб-приложению

Веб-приложение должно предоставлять пользователю следующие возможности:

- просмотр сводки на Dashboard;
- загрузка CSV-файлов с выбором полигона и указанием загрузившего;
- просмотр таблицы измерений;
- фильтрация измерений по полигону, типу показателя, периоду и файлу импорта;
- ручное создание, редактирование и удаление измерений;
- отображение источника записи: имя CSV-файла или «Введено вручную»;
- построение графика нескольких показателей на одном полигоне;
- сравнение одного показателя на нескольких полигонах;
- создание, редактирование и удаление полигонов;
- создание, редактирование и удаление типов показателей;
- создание, редактирование и удаление единиц измерения;
- вывод понятного сообщения, если удаление справочной записи невозможно из-за связанных данных.

# 2. Проектирование базы данных

## 2.1. Выделение основных сущностей предметной области

Логическая модель построена вокруг факта экологического измерения. В рамках проектирования выделены шесть таблиц: polygons, data_collectors, measurement_units, sensor_types, import_files и measurements.

Таблица polygons хранит полигоны наблюдения. Таблица data_collectors хранит фамилии людей, загружавших CSV-файлы. Таблица measurement_units хранит единицы измерения. Таблица sensor_types хранит типы измеряемых показателей и связывает каждый тип с единицей измерения. Таблица import_files фиксирует факт загрузки CSV-файла. Таблица measurements хранит нормализованные значения измерений.

## 2.2. Описание таблиц базы данных

| Таблица | Назначение |
| --- | --- |
| polygons | Справочник экологических полигонов. |
| data_collectors | Справочник загрузивших CSV-файлы сотрудников. |
| measurement_units | Справочник единиц измерения. |
| sensor_types | Справочник типов измеряемых показателей; каждый тип связан с единицей измерения. |
| import_files | Журнал CSV-загрузок: файл, полигон, загрузивший, статус и количество обработанных строк. |
| measurements | Основная таблица фактов измерений: полигон, тип показателя, единица, время, значение и источник импорта. |

Ниже приведено описание атрибутов таблиц базы данных. Типы данных указаны с учетом выбранной СУБД MySQL 8.

Поля таблицы polygons:

| Поле | Тип данных | Назначение |
| --- | --- | --- |
| polygon_id | BIGINT, PK, AUTO_INCREMENT | Идентификатор полигона. |
| name | VARCHAR(255), NOT NULL | Название полигона. |
| location | VARCHAR(255), NULL | Текстовое описание местоположения. |
| description | TEXT, NULL | Дополнительное описание полигона. |
| created_at | DATETIME, NOT NULL | Дата и время создания записи. |
| updated_at | DATETIME, NULL | Дата и время последнего изменения. |

Поля таблицы data_collectors:

| Поле | Тип данных | Назначение |
| --- | --- | --- |
| collector_id | BIGINT, PK, AUTO_INCREMENT | Идентификатор загрузившего данные сотрудника. |
| last_name | VARCHAR(255), NOT NULL, INDEX | Фамилия загрузившего; используется при поиске существующей записи. |
| first_name | VARCHAR(255), NULL | Имя загрузившего, если указано при импорте. |
| middle_name | VARCHAR(255), NULL | Отчество загрузившего, если указано при импорте. |
| created_at | DATETIME, NOT NULL | Дата и время создания записи. |

Поля таблицы measurement_units:

| Поле | Тип данных | Назначение |
| --- | --- | --- |
| unit_id | BIGINT, PK, AUTO_INCREMENT | Идентификатор единицы измерения. |
| name | VARCHAR(255), NOT NULL | Полное название единицы измерения. |
| symbol | VARCHAR(50), NOT NULL, UNIQUE | Краткое обозначение единицы, например %, ppm, hPa. |
| description | TEXT, NULL | Описание области применения единицы. |

Поля таблицы sensor_types:

| Поле | Тип данных | Назначение |
| --- | --- | --- |
| sensor_type_id | BIGINT, PK, AUTO_INCREMENT | Идентификатор типа измеряемого показателя. |
| unit_id | BIGINT, FK, NOT NULL | Ссылка на единицу измерения из measurement_units. |
| name | VARCHAR(255), NOT NULL | Название показателя, отображаемое в интерфейсе. |
| code | VARCHAR(100), NOT NULL, UNIQUE | Системный код показателя для сопоставления с CSV-колонками. |
| description | TEXT, NULL | Дополнительное описание показателя. |

Поля таблицы import_files:

| Поле | Тип данных | Назначение |
| --- | --- | --- |
| import_file_id | BIGINT, PK, AUTO_INCREMENT | Идентификатор записи журнала импорта. |
| polygon_id | BIGINT, FK, NOT NULL | Полигон, к которому относится загруженный CSV-файл. |
| collector_id | BIGINT, FK, NOT NULL | Сотрудник, выполнивший загрузку. |
| file_name | VARCHAR(255), NOT NULL | Имя загруженного CSV-файла. |
| uploaded_at | DATETIME, NOT NULL | Дата и время загрузки файла. |
| rows_count | INT, NOT NULL | Количество обработанных строк CSV. |
| measurements_count | INT, NOT NULL | Количество созданных измерений. |
| status | VARCHAR(50), NOT NULL | Статус обработки файла. |
| error_message | TEXT, NULL | Текст ошибки, если импорт завершился неуспешно. |

Поля таблицы measurements:

| Поле | Тип данных | Назначение |
| --- | --- | --- |
| measurement_id | BIGINT, PK, AUTO_INCREMENT | Идентификатор измерения. |
| polygon_id | BIGINT, FK, NOT NULL | Полигон, на котором получено значение. |
| sensor_type_id | BIGINT, FK, NOT NULL | Тип измеряемого показателя. |
| unit_id | BIGINT, FK, NOT NULL | Единица измерения значения. |
| import_file_id | BIGINT, FK, NULL | CSV-файл-источник; для ручных записей равно NULL. |
| measured_at | DATETIME, NOT NULL | Дата и время измерения. |
| value | DECIMAL(16,4), NOT NULL | Числовое значение показателя. |
| created_at | DATETIME, NOT NULL | Дата и время создания записи в базе данных. |

Для физического хранения выбрана СУБД MySQL 8. Таблицы рассчитаны на использование транзакционного механизма InnoDB, а идентификаторы сущностей задаются как BIGINT с автоинкрементом. Для строковых атрибутов используются VARCHAR, для расширенных описаний — TEXT, для дат и времени — DATETIME, для числовых значений измерений — DECIMAL(16,4).

Исходная широкая форма CSV используется только как формат ввода данных. При загрузке значения из столбцов показателей преобразуются в строки таблицы measurements: одна непустая ячейка показателя соответствует одному факту измерения. Такой подход соответствует реляционной модели, уменьшает зависимость схемы от состава CSV-столбцов и упрощает фильтрацию, агрегацию и построение графиков.

Та же таблица measurements используется для ручных измерений. Для записей, полученных из CSV, поле import_file_id содержит ссылку на import_files; для ручных записей это поле остается пустым. Благодаря этому источник данных фиксируется без добавления дополнительных сущностей.

## 2.3. ER-диаграмма базы данных

![ER-диаграмма базы данных экологических измерений](assets/fig01_er_diagram.png)

Рисунок 1 — ER-диаграмма базы данных экологических измерений

ER-диаграмма показывает таблицы, поля, типы данных, первичные ключи, внешние ключи, уникальные ограничения и основные связи. Центральной таблицей является measurements. Она связана с polygons, sensor_types, measurement_units и, при файловом импорте, с import_files. Таблица import_files связана с polygons и data_collectors. Таблица sensor_types связана с measurement_units.

ER-модель предусматривает два варианта происхождения измерений: файловый импорт и ручной ввод. Оба варианта используют одну структуру таблицы measurements. Различие фиксируется значением import_file_id: при импорте хранится ссылка на запись журнала import_files, а при ручном вводе связь с журналом отсутствует.

## 2.4. Нормализация базы данных до третьей нормальной формы

Первая нормальная форма обеспечивается тем, что каждое поле содержит атомарное значение, а повторяющиеся группы CSV-столбцов преобразуются в строки measurements. Вторая нормальная форма соблюдается за счет простых первичных ключей: неключевые атрибуты зависят от ключа своей таблицы. Третья нормальная форма соблюдается, потому что названия полигонов, типов показателей и единиц измерения вынесены в справочники и не дублируются в таблице измерений.

Поле unit_id в measurements хранится как ссылка на примененную единицу измерения. При ручном создании записи оно не вводится пользователем произвольно, а определяется по sensor_types.unit_id. Это предотвращает противоречивые комбинации данных и сохраняет целостность справочников.

## 2.5. Индексы для ускорения фильтрации и построения графиков

Таблица measurements является наиболее крупной таблицей, поэтому в модели заданы составные индексы для основных сценариев выборки. Индекс idx_measurements_polygon_sensor_time используется при фильтрации по полигону, типу показателя и периоду. Индекс idx_measurements_sensor_time_polygon полезен при сравнении одного показателя на нескольких полигонах. Индекс idx_measurements_time ускоряет выборки по периоду, а idx_measurements_import_file используется при фильтрации по файлу импорта.

Для журнала загрузок задан индекс idx_import_files_polygon_uploaded, который ускоряет просмотр импортов по полигону и дате загрузки. Для справочника загрузивших используется индекс по фамилии, так как при импорте система ищет существующую запись по last_name.

# 3. Серверная часть и работа с базой данных

## 3.1. Общая роль серверной части

Серверная часть реализована на FastAPI. Она принимает HTTP-запросы от frontend, проверяет входные данные через Pydantic-схемы, выполняет операции через SQLAlchemy ORM и возвращает JSON-ответы. Основные группы маршрутов: справочники, импорт CSV, измерения, графики, Dashboard и health-check.

Слой сервисов содержит бизнес-логику работы с базой данных. Например, сервис ручного добавления измерения проверяет существование полигона и типа показателя, получает единицу измерения из выбранного типа и сохраняет запись. Сервис справочников перед удалением считает связанные записи и блокирует операцию, если удаление нарушило бы целостность.

## 3.2. Подключение backend к MySQL

Подключение к MySQL формируется в backend/app/core/config.py, а engine и сессии создаются в backend/app/database.py.

Листинг 1 — Подключение backend к базе данных MySQL

```python
@property
def sqlalchemy_database_uri(self) -> str:
    return (
        f"mysql+pymysql://{self.db_user}:{self.db_password}"
        f"@{self.db_host}:{self.db_port}/{self.db_name}?charset=utf8mb4"
    )

engine = create_engine(
    settings.sqlalchemy_database_uri,
    pool_pre_ping=True,
    future=True,
)

SessionLocal = sessionmaker(
    autocommit=False,
    autoflush=False,
    bind=engine,
    class_=Session,
)
```

В коде создается SQLAlchemy engine с драйвером mysql+pymysql. Параметр pool_pre_ping=True проверяет соединение перед использованием, что важно для сетевой базы данных и контейнерной среды.

## 3.3. ORM-запросы и операции изменения данных

Ниже приведены короткие фрагменты реального кода, связанные только с обращением к базе данных. Полные файлы не приводятся, чтобы не перегружать отчет.

Листинг 2 — ORM-запрос для получения таблицы измерений

```python
base_stmt = (
    select(
        Measurement.measurement_id.label("measurement_id"),
        Measurement.measured_at.label("measured_at"),
        Polygon.name.label("polygon_name"),
        SensorType.name.label("sensor_name"),
        Measurement.value.label("value"),
        MeasurementUnit.symbol.label("unit_symbol"),
        import_alias.file_name.label("file_name"),
        collector_alias.last_name.label("collector_last_name"),
    )
    .join(Polygon, Polygon.polygon_id == Measurement.polygon_id)
    .join(SensorType, SensorType.sensor_type_id == Measurement.sensor_type_id)
    .join(MeasurementUnit, MeasurementUnit.unit_id == Measurement.unit_id)
    .outerjoin(import_alias, import_alias.import_file_id == Measurement.import_file_id)
    .outerjoin(collector_alias, collector_alias.collector_id == import_alias.collector_id)
)
```

Запрос формирует строки таблицы измерений вместе с данными справочников. Внешнее соединение с import_files требуется из-за ручных записей, у которых import_file_id равен NULL.

Листинг 3 — ORM-операция ручного добавления измерения

```python
sensor_type = db.get(SensorType, payload.sensor_type_id)

measurement = Measurement(
    polygon_id=payload.polygon_id,
    sensor_type_id=payload.sensor_type_id,
    unit_id=sensor_type.unit_id,
    import_file_id=None,
    measured_at=payload.measured_at,
    value=payload.value,
)
db.add(measurement)
db.commit()
db.refresh(measurement)
```

При ручном вводе связь с файлом импорта не создается. Единица измерения берется из выбранного типа показателя, а не вводится пользователем отдельно.

Листинг 4 — ORM-операция редактирования измерения

```python
measurement.polygon_id = payload.polygon_id
measurement.sensor_type_id = payload.sensor_type_id
measurement.unit_id = sensor_type.unit_id
measurement.measured_at = payload.measured_at
measurement.value = payload.value

db.commit()
db.refresh(measurement)
```

Редактирование меняет основные поля измерения, но не очищает import_file_id. Поэтому импортированная запись сохраняет связь с исходным CSV-файлом.

Листинг 5 — Проверка связей перед удалением полигона

```python
has_measurements = db.scalar(
    select(func.count(Measurement.measurement_id))
    .where(Measurement.polygon_id == polygon_id)
) or 0

has_imports = db.scalar(
    select(func.count(ImportFile.import_file_id))
    .where(ImportFile.polygon_id == polygon_id)
) or 0

if has_measurements or has_imports:
    raise HTTPException(status_code=409, detail="Полигон используется.")
```

Фрагмент показывает фактический подход к защите справочников: перед удалением выполняется проверка связанных измерений и загрузок.

Листинг 6 — ORM-запрос для построения графика

```python
stmt = (
    select(
        SensorType.sensor_type_id.label("sensor_type_id"),
        SensorType.name.label("sensor_name"),
        bucket,
        value_expr.label("value"),
    )
    .join(SensorType, SensorType.sensor_type_id == Measurement.sensor_type_id)
    .where(and_(
        Measurement.polygon_id == polygon_id,
        Measurement.sensor_type_id.in_(sensor_ids),
    ))
)

rows = db.execute(stmt).mappings().all()
```

Запрос строится по таблице measurements, фильтрует данные по полигону и списку типов показателей и возвращает точки временных рядов для графика.

# 4. Пользовательский интерфейс веб-приложения

## 4.1. Главная страница Dashboard

Dashboard показывает общее состояние системы: количество полигонов, типов показателей, измерений и CSV-загрузок. На странице также размещены быстрые переходы к загрузке CSV, таблице измерений, графикам и справочникам. Ниже отображается блок последних загрузок.

![Главная страница Dashboard](assets/fig02_dashboard.png)

Рисунок 2 — Главная страница Dashboard

## 4.2. Страница загрузки CSV-файлов

Страница загрузки содержит выбор полигона, поля для фамилии, имени и отчества загрузившего, поле выбора CSV-файла, подсказку по формату и кнопку отправки. Пользователь видит ожидаемую структуру CSV: обязательный столбец даты и столбцы показателей.

![Страница загрузки CSV-файла](assets/fig03_upload_page.png)

Рисунок 3 — Страница загрузки CSV-файла

## 4.3. Результат загрузки данных

После успешной загрузки на странице появляется блок результата. В нем отображаются имя файла, статус обработки, количество строк CSV, количество созданных измерений и количество пропущенных значений. Также есть кнопки перехода к измерениям и графикам.

![Результат загрузки данных](assets/fig04_upload_result.png)

Рисунок 4 — Результат загрузки данных

## 4.4. Просмотр и ручное управление измерениями

Страница измерений содержит таблицу с колонками: дата и время, полигон, тип показателя, значение, единица измерения, источник, фамилия и действия. В колонке источника отображается имя CSV-файла или значение «Введено вручную». Для каждой строки доступны кнопки редактирования и удаления.

![Таблица экологических измерений с операциями управления записями](assets/fig05_measurements_table.png)

Рисунок 5 — Таблица экологических измерений с операциями управления записями

Форма ручного добавления и редактирования расположена на той же странице. В ней выбираются полигон и тип показателя, вводятся дата, время и значение. Поле фамилии используется при редактировании импортированной записи; ручная запись не требует файла-источника.

![Форма ручного добавления и редактирования измерения](assets/fig06_measurement_form.png)

Рисунок 6 — Форма ручного добавления и редактирования измерения

## 4.5. Фильтрация данных по полигону, датчику и периоду

Блок фильтров позволяет выбрать полигон, тип показателя, файл импорта, начальную и конечную дату, а также порядок сортировки. После применения фильтра таблица показывает только записи, соответствующие выбранным условиям.

![Фильтрация данных](assets/fig07_measurements_filters.png)

Рисунок 7 — Фильтрация данных

## 4.6. Построение графика нескольких датчиков на одном полигоне

На странице графиков первый режим предназначен для анализа нескольких показателей на одном полигоне. Пользователь выбирает полигон, несколько типов показателей, период и режим агрегации. График отображает отдельную линию для каждого выбранного показателя.

![График нескольких датчиков на одном полигоне](assets/fig08_chart_multi_sensor.png)

Рисунок 8 — График нескольких датчиков на одном полигоне

## 4.7. Сравнение одного датчика на нескольких полигонах

Второй режим графиков используется для сравнения одного показателя на нескольких полигонах. Пользователь выбирает тип показателя, несколько полигонов, период и агрегацию. Каждая линия соответствует отдельному полигону.

![Сравнение показателя на нескольких полигонах](assets/fig09_chart_multi_polygon.png)

Рисунок 9 — Сравнение показателя на нескольких полигонах

## 4.8. Управление справочниками

Страница справочников содержит вкладки для полигонов, типов показателей и единиц измерения. На каждой вкладке есть форма создания или редактирования записи, таблица существующих записей и кнопки действий. Для полигонов отображаются название, локация и описание. Для типов показателей отображаются название, код, единица измерения и описание. Для единиц измерения отображаются название, символ и описание.

![Страница управления справочниками](assets/fig10_references_page.png)

Рисунок 10 — Страница управления справочниками

Форма справочника используется как для добавления, так и для редактирования. При выборе действия редактирования поля формы заполняются данными выбранной строки, после чего пользователь может сохранить изменения.

![Добавление и редактирование записи справочника](assets/fig11_reference_form.png)

Рисунок 11 — Добавление и редактирование записи справочника

## 4.9. Журнал загруженных файлов

История загрузок видна на Dashboard в блоке последних импортов. В строках журнала отображаются имя файла, полигон, фамилия загрузившего, дата загрузки, статус и количество созданных измерений. Ручные измерения в этом блоке не появляются, потому что они не создаются из CSV-файлов.

![Журнал загруженных файлов](assets/fig12_import_journal.png)

Рисунок 12 — Журнал загруженных файлов

# 5. SQL-запросы к базе данных

## 5.1. Запрос для вывода измерений с данными из справочников

Назначение запроса — получить измерения вместе с названиями полигона, типа показателя, единицей измерения и сведениями о файле импорта.

```sql
SELECT
    m.measurement_id,
    m.measured_at,
    p.name AS polygon_name,
    st.name AS sensor_type_name,
    m.value,
    mu.symbol AS unit_symbol,
    f.file_name,
    dc.last_name AS collector_last_name
FROM measurements AS m
JOIN polygons AS p ON p.polygon_id = m.polygon_id
JOIN sensor_types AS st ON st.sensor_type_id = m.sensor_type_id
JOIN measurement_units AS mu ON mu.unit_id = m.unit_id
LEFT JOIN import_files AS f ON f.import_file_id = m.import_file_id
LEFT JOIN data_collectors AS dc ON dc.collector_id = f.collector_id
ORDER BY m.measured_at DESC
LIMIT 50;
```

Результат используется для таблицы измерений. LEFT JOIN нужен, потому что ручные записи не связаны с import_files.

## 5.2. Запрос для фильтрации измерений по полигону и датчику

```sql
SELECT
    m.measured_at,
    p.name AS polygon_name,
    st.name AS sensor_type_name,
    m.value,
    mu.symbol AS unit_symbol
FROM measurements AS m
JOIN polygons AS p ON p.polygon_id = m.polygon_id
JOIN sensor_types AS st ON st.sensor_type_id = m.sensor_type_id
JOIN measurement_units AS mu ON mu.unit_id = m.unit_id
WHERE m.polygon_id = :polygon_id
  AND m.sensor_type_id = :sensor_type_id
  AND m.measured_at BETWEEN :date_from AND :date_to
ORDER BY m.measured_at ASC;
```

Запрос возвращает значения одного показателя на одном полигоне за выбранный период.

## 5.3. Запрос для построения графика нескольких датчиков на одном полигоне

```sql
SELECT
    st.sensor_type_id,
    st.name AS sensor_type_name,
    mu.symbol AS unit_symbol,
    m.measured_at,
    m.value
FROM measurements AS m
JOIN sensor_types AS st ON st.sensor_type_id = m.sensor_type_id
JOIN measurement_units AS mu ON mu.unit_id = m.unit_id
WHERE m.polygon_id = :polygon_id
  AND m.sensor_type_id IN (:sensor_type_ids)
  AND m.measured_at BETWEEN :date_from AND :date_to
ORDER BY m.measured_at ASC, st.sensor_type_id ASC;
```

Результат группируется по sensor_type_id, после чего каждая группа отображается как отдельная серия графика.

## 5.4. Запрос для сравнения одного датчика на нескольких полигонах

```sql
SELECT
    p.polygon_id,
    p.name AS polygon_name,
    m.measured_at,
    m.value,
    mu.symbol AS unit_symbol
FROM measurements AS m
JOIN polygons AS p ON p.polygon_id = m.polygon_id
JOIN measurement_units AS mu ON mu.unit_id = m.unit_id
WHERE m.sensor_type_id = :sensor_type_id
  AND m.polygon_id IN (:polygon_ids)
  AND m.measured_at BETWEEN :date_from AND :date_to
ORDER BY m.measured_at ASC, p.polygon_id ASC;
```

Запрос возвращает временные ряды одного показателя для нескольких полигонов.

## 5.5. Запрос для просмотра истории загрузок CSV-файлов

```sql
SELECT
    f.import_file_id,
    f.file_name,
    p.name AS polygon_name,
    dc.last_name AS collector_last_name,
    f.uploaded_at,
    f.status,
    f.rows_count,
    f.measurements_count,
    f.error_message
FROM import_files AS f
JOIN polygons AS p ON p.polygon_id = f.polygon_id
JOIN data_collectors AS dc ON dc.collector_id = f.collector_id
ORDER BY f.uploaded_at DESC;
```

Запрос показывает историю файлового импорта. Ручные измерения в этом результате отсутствуют.

## 5.6. Запрос для агрегации измерений по времени

```sql
SELECT
    st.name AS sensor_type_name,
    DATE_FORMAT(m.measured_at, '%Y-%m-%d %H:00:00') AS time_bucket,
    AVG(m.value) AS avg_value,
    mu.symbol AS unit_symbol
FROM measurements AS m
JOIN sensor_types AS st ON st.sensor_type_id = m.sensor_type_id
JOIN measurement_units AS mu ON mu.unit_id = m.unit_id
WHERE m.polygon_id = :polygon_id
  AND m.sensor_type_id IN (:sensor_type_ids)
  AND m.measured_at BETWEEN :date_from AND :date_to
GROUP BY st.name, mu.symbol, time_bucket
ORDER BY time_bucket ASC;
```

Запрос используется для почасовой агрегации. Для дневной агрегации формат временного интервала изменяется на день.

## 5.7. Запросы добавления, редактирования и удаления записей

В программной реализации эти операции выполняются через SQLAlchemy ORM. Ниже приведены SQL-эквиваленты действий на уровне базы данных.

Добавление полигона:

```sql
INSERT INTO polygons (name, location, description)
VALUES ('Полигон №4', 'Восточная зона', 'Дополнительная точка наблюдения');
```

Редактирование типа показателя:

```sql
UPDATE sensor_types
SET name = 'Температура воздуха', description = 'Температура окружающей среды'
WHERE sensor_type_id = :sensor_type_id;
```

Ручное добавление измерения. В реальном backend unit_id определяется по выбранному типу показателя.

```sql
INSERT INTO measurements (
    polygon_id,
    sensor_type_id,
    unit_id,
    import_file_id,
    measured_at,
    value
)
VALUES (
    :polygon_id,
    :sensor_type_id,
    (SELECT unit_id FROM sensor_types WHERE sensor_type_id = :sensor_type_id),
    NULL,
    :measured_at,
    :value
);
```

Редактирование измерения:

```sql
UPDATE measurements
SET polygon_id = :polygon_id,
    sensor_type_id = :sensor_type_id,
    unit_id = (SELECT unit_id FROM sensor_types WHERE sensor_type_id = :sensor_type_id),
    measured_at = :measured_at,
    value = :value
WHERE measurement_id = :measurement_id;
```

Удаление измерения:

```sql
DELETE FROM measurements
WHERE measurement_id = :measurement_id;
```

Проверка возможности удаления полигона:

```sql
SELECT
    (SELECT COUNT(*) FROM measurements WHERE polygon_id = :polygon_id) AS measurements_count,
    (SELECT COUNT(*) FROM import_files WHERE polygon_id = :polygon_id) AS imports_count;
```

Если хотя бы один счетчик больше нуля, полигон удалять нельзя. Аналогично тип показателя проверяется по measurements, а единица измерения — по sensor_types и measurements.

# Заключение

В ходе курсовой работы была разработана база данных MySQL и веб-приложение для хранения и анализа экологических измерений. Цель работы достигнута: данные из CSV-файлов преобразуются в нормализованный длинный формат и сохраняются в таблице measurements, а пользователь может работать с ними через интерфейс приложения.

Спроектирована структура из таблиц polygons, data_collectors, measurement_units, sensor_types, import_files и measurements. ER-диаграмма отражает первичные ключи, внешние ключи, типы данных, уникальные ограничения и nullable-связь для ручных измерений. Структура соответствует третьей нормальной форме, поскольку справочные данные не дублируются в таблице измерений.

Реализованы импорт CSV, просмотр измерений, фильтрация, построение графиков, сравнение полигонов, ручное добавление, редактирование и удаление измерений, а также управление справочниками полигонов, типов показателей и единиц измерения. Импортированные и ручные записи разделяются по признаку import_file_id: у импортированных записей есть ссылка на файл, у ручных записей это поле равно NULL.

Разработанная система демонстрирует работу с сетевой базой данных через backend API и ORM-слой. Пользователь выполняет операции с MySQL через интерфейс, без ручного написания SQL-команд. Контроль ссылочной целостности предотвращает удаление справочных записей, которые уже используются существующими данными.

# Список использованных источников

1. MySQL 8.0 Reference Manual [Электронный ресурс]. — URL: https://dev.mysql.com/doc/refman/8.0/en/.
2. FastAPI Documentation [Электронный ресурс]. — URL: https://fastapi.tiangolo.com/.
3. SQLAlchemy 2.0 Documentation [Электронный ресурс]. — URL: https://docs.sqlalchemy.org/en/20/.
4. Alembic Documentation [Электронный ресурс]. — URL: https://alembic.sqlalchemy.org/.
5. Angular Documentation [Электронный ресурс]. — URL: https://angular.dev/.
6. Apache ECharts Documentation [Электронный ресурс]. — URL: https://echarts.apache.org/en/index.html.
7. Docker Documentation [Электронный ресурс]. — URL: https://docs.docker.com/.
8. Дейт К. Дж. Введение в системы баз данных. — 8-е изд. — Москва: Вильямс, 2005.
9. Силбершац А., Корт Г., Сударшан С. Основы баз данных. — Москва: Вильямс, 2011.
10. Материалы проекта eco-monitoring-app: README.md, PROJECT_REQUIREMENTS.md, исходный код backend и frontend.
"""


def ensure_markdown() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    if not MD_PATH.exists():
        MD_PATH.write_text(DEFAULT_MARKDOWN, encoding="utf-8")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name="Times New Roman", size=14, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    set_run_font(run, size=12)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for style_name, size, before, after, align in (
        ("Heading 1", 14, 6, 3, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, 6, 2, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, 4, 2, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = align
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)

    return doc


def add_heading(doc: Document, text: str, level: int, first_top_heading_seen: bool) -> bool:
    # Do not force every chapter onto a new page: the target report size is
    # about 30-35 pages and the user will add a separate title page later.
    paragraph = doc.add_heading(text, level=min(level, 3))
    for run in paragraph.runs:
        set_run_font(run, size=14, bold=True)
    return True if level == 1 else first_top_heading_seen


def add_body_paragraph(doc: Document, text: str, *, no_indent: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(0 if no_indent else 1.25)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=14)


def add_markdown_image(doc: Document, target: str) -> None:
    image_path = (BASE_DIR / target).resolve()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    # A4 usable width with the configured margins is about 16.5 cm.
    run.add_picture(str(image_path), width=Cm(16.0))


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=14)


def add_placeholder(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    shade_paragraph(paragraph, "F2F2F2")
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True, italic=True)


def add_list_item(doc: Document, text: str, ordered: bool = False) -> None:
    paragraph = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=14)


def add_code_block(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").splitlines():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.25)
        paragraph.paragraph_format.right_indent = Cm(0.1)
        paragraph.paragraph_format.line_spacing = Pt(11)
        paragraph.paragraph_format.space_after = Pt(0)
        shade_paragraph(paragraph, "F2F2F2")
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, name="Consolas", size=10)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.first_line_indent = Cm(0)


def table_widths(column_count: int) -> list[float]:
    if column_count == 2:
        return [5.0, 11.5]
    if column_count == 3:
        return [3.5, 3.5, 9.5]
    if column_count == 4:
        return [3.0, 3.5, 5.0, 5.0]
    if column_count == 5:
        return [2.5, 3.0, 3.0, 4.0, 4.0]
    return [16.5 / max(column_count, 1)] * column_count


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    widths = table_widths(column_count)

    header_cells = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        cell = header_cells[i]
        cell.width = Cm(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(value)
        set_run_font(run, size=9.5, bold=True)

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cell = cells[i]
            cell.width = Cm(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            set_run_font(run, size=9.5)

    after = doc.add_paragraph()
    after.paragraph_format.first_line_indent = Cm(0)
    after.paragraph_format.space_after = Pt(4)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        stripped = lines[i].strip()
        if not (stripped.startswith("|") and stripped.endswith("|")):
            break
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build_docx_from_markdown(markdown: str) -> Document:
    doc = configure_document()
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    top_heading_seen = False
    current_top_heading = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lines = []
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, new_i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            i = new_i
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1:
                current_top_heading = text
            top_heading_seen = add_heading(doc, text, level, top_heading_seen)
            i += 1
            continue

        image_match = re.match(r"^!\[[^\]]*\]\(([^)]+)\)$", stripped)
        if image_match:
            add_markdown_image(doc, image_match.group(1).strip())
            i += 1
            continue

        if stripped.startswith("[Место для рисунка"):
            add_placeholder(doc, stripped)
            i += 1
            continue

        if stripped.startswith(("Рисунок ", "Листинг ", "Таблица ")):
            add_caption(doc, stripped)
            i += 1
            continue

        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:].strip(), ordered=False)
            i += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match and current_top_heading == "Список использованных источников":
            add_list_item(doc, ordered_match.group(1).strip(), ordered=True)
            i += 1
            continue

        add_body_paragraph(doc, stripped, no_indent=current_top_heading == "Содержание")
        i += 1

    return doc


def main() -> None:
    ensure_markdown()
    markdown = MD_PATH.read_text(encoding="utf-8-sig")
    doc = build_docx_from_markdown(markdown)
    doc.save(DOCX_PATH)
    print(f"Markdown: {MD_PATH}")
    print(f"DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()


